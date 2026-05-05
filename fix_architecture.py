"""Replace the ASCII pipeline flow in section 5 with a clean table."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = "Project Report.docx"
OUT = "Project Report.docx"


def set_table_borders(tbl):
    tblPr = tbl._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._element.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)


PIPELINE_ROWS = [
    ["Stage", "Module", "What it does"],
    ["1. Input", "(file dialog)", "User opens a JPG or PNG image"],
    ["2. Preprocess", "preprocessing.py",
        "Upscale, denoise, CLAHE contrast, mild sharpening; "
        "optional red-ink channel for headers"],
    ["3. OCR", "ocr_engine.py",
        "Extract text using EasyOCR with Tesseract as a fallback; "
        "returns text and per-word bounding boxes"],
    ["4. Detect formatting", "formatting_detector.py",
        "Group words into lines and paragraphs; classify "
        "headings, bullets, alignment, and two-column layout"],
    ["5. Generate document", "docx_generator.py",
        "Build a styled Word document from the FormattedBlocks"],
    ["6. Output", ".docx file",
        "User clicks Save; the editable text is written to disk"],
]


def main():
    doc = Document(SRC)

    # Locate the ASCII flow paragraphs by their distinctive content
    paragraphs = list(doc.paragraphs)
    start = None
    end = None
    for i, p in enumerate(paragraphs):
        txt = p.text.strip()
        if start is None and txt == "Input Image (JPG/PNG)":
            start = i
        if start is not None and txt == "Output (.docx file)":
            end = i
            break

    if start is None or end is None:
        raise RuntimeError("Could not locate the ASCII flowchart paragraphs")

    print(f"Removing paragraphs {start} through {end} (the ASCII flow)")

    # Capture the element that comes RIGHT AFTER the flow (so we can insert
    # the new table there). If the next sibling is the GUI sentence, the new
    # table will appear in the same place.
    after_flow_xml = paragraphs[end]._element.getnext()
    if after_flow_xml is None:
        raise RuntimeError("No paragraph after the ASCII flow — unexpected")

    # Remove the ASCII paragraphs
    body = doc.element.body
    for i in range(end, start - 1, -1):
        body.remove(paragraphs[i]._element)

    # Insert an intro paragraph and a table just before after_flow_xml
    # 1) intro text
    intro = doc.add_paragraph()
    intro.add_run(
        "The pipeline is summarised below. Each stage takes the output "
        "of the previous one as input and passes its result to the next."
    )
    intro_xml = intro._element

    # 2) build the table at end then move it
    tbl = doc.add_table(rows=len(PIPELINE_ROWS), cols=len(PIPELINE_ROWS[0]))
    set_table_borders(tbl)
    for ri, row in enumerate(PIPELINE_ROWS):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            if ri == 0:
                run.bold = True
    tbl_xml = tbl._element

    # 3) blank paragraph after the table
    spacer = doc.add_paragraph()
    spacer_xml = spacer._element

    # Move all three before after_flow_xml in order: intro, tbl, spacer
    after_flow_xml.addprevious(intro_xml)
    after_flow_xml.addprevious(tbl_xml)
    after_flow_xml.addprevious(spacer_xml)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
